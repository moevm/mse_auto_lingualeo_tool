from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def save_to_file(wordset_name, words, file_name=None, many_translations=False):
    document = Document()
    document.add_heading(wordset_name, 0)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    if file_name is None:
        file_name = wordset_name

    header_content = {
        0: 'Слово на английском',
        1: 'Транскрипция',
        2: 'Перевод'
    }

    hdr_cells = table.rows[0].cells

    for idx, cell in enumerate(hdr_cells):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = header_content[idx]

    for word in words:
        row_cells = table.add_row().cells
        row_cells[0].text = word['wd']
        row_cells[1].text = word['scr'] if word['scr'] != None else '-'

        if many_translations:
            translations = map(lambda word: word['tr'], word['trs'])
            deduplicated_translations = list(dict.fromkeys(translations))
            result = ", ".join(deduplicated_translations)
            row_cells[2].text = result
        else:
            row_cells[2].text = word['trs'][0]['tr']

    document.save('%s.docx' % file_name)
