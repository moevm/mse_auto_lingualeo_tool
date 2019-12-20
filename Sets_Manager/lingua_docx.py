from docx import Document


# Класс парсера docx файла, хранящего в себе набор слов для серсива Lingualeo
class LinguaDocxParser:
    # Конструктор парсера от пути к файлу
    def __init__(self, file_name):
        self.docx_document = Document(file_name)


    # Функция печати содержимого docx файла в консоль
    def print_file(self):
        for paragraph in self.docx_document.paragraphs:
            print(paragraph.text)


    # Функция, создающая список слов из docx файла вместе с переводами и транскрипцией
    def create_word_set(self):
        word_set, dublicates = [], []
        for paragraph in self.docx_document.paragraphs:
            word_set_elem_parts = paragraph.text.split('—')
            word_set_elem = dict.fromkeys(["word", "transcription", "translation"])
            if word_set_elem_parts[0]:
                if len(word_set_elem_parts) == 3:
                    word_set_elem["word"] = word_set_elem_parts[0].strip()
                    word_set_elem["transcription"] = word_set_elem_parts[1].strip()
                    word_set_elem["translation"] = word_set_elem_parts[2].strip()
                elif len(word_set_elem_parts) == 2:
                    word_set_elem["word"] = word_set_elem_parts[0].strip()
                    word_set_elem["translation"] = word_set_elem_parts[1].strip()
                else:
                    word_set_elem["word"] = word_set_elem_parts[0].strip()
                already_exists = False
                for elem in word_set:
                    if elem["word"] == word_set_elem["word"]:
                        already_exists = True
                        dublicates.append(word_set_elem)
                        break
                if not already_exists:
                    word_set.append(word_set_elem)
        for dublicated_elem in dublicates:
            if dublicated_elem["translation"] is None:
                continue
            for word_set_elem in word_set:
                if dublicated_elem["word"] == word_set_elem["word"]:
                    if word_set_elem["translation"] is None:
                        word_set_elem["translation"] = dublicated_elem["translation"]
                    elif dublicated_elem["translation"] != word_set_elem["translation"] and \
                            word_set_elem["translation"].find(dublicated_elem["translation"]) == -1:
                        word_set_elem["translation"] += '; ' + dublicated_elem["translation"] 
        return word_set

